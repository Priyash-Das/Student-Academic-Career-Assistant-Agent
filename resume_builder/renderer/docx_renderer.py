from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from resume_builder.renderer.layout_constants import (
    FONT_NAME,
    FONT_SIZE_HEADER_NAME,
    FONT_SIZE_HEADER_CONTACT,
    FONT_SIZE_SECTION_TITLE,
    FONT_SIZE_SECTION_CONTENT,
    FONT_SIZE_BULLET,
    PAGE_MARGIN_INCH,
    LINE_SPACING,
    SECTION_SPACING,
    BULLET_INDENT,
)
from resume_builder.core.resume_schema import ResumeSchema
class DocxRenderer:
    @staticmethod
    def render(resume: ResumeSchema, output_path: str):
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(PAGE_MARGIN_INCH)
        section.bottom_margin = Inches(PAGE_MARGIN_INCH)
        section.left_margin = Inches(PAGE_MARGIN_INCH)
        section.right_margin = Inches(PAGE_MARGIN_INCH)
        DocxRenderer._setup_styles(doc)
        DocxRenderer._render_header(doc, resume.header)
        doc.add_paragraph()
        if resume.summary:
            DocxRenderer._render_section_title(doc, "PROFESSIONAL SUMMARY")
            p = doc.add_paragraph(resume.summary)
            p.style = "ResumeNormal"
            p.paragraph_format.space_after = Pt(SECTION_SPACING)
        if resume.experience:
            DocxRenderer._render_section_title(doc, "EXPERIENCE")
            for exp in resume.experience:
                role_line = f"{exp.role} – {exp.company}"
                if exp.location:
                    role_line += f", {exp.location}"
                p = doc.add_paragraph()
                p.style = "ResumeHeading"
                left_run = p.add_run(role_line)
                left_run.bold = True
                left_run.font.name = FONT_NAME
                left_run.font.size = Pt(FONT_SIZE_SECTION_CONTENT)
                if exp.dates:
                    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.0))
                    p.add_run("\t").bold = False
                    right_run = p.add_run(exp.dates)
                    right_run.font.name = FONT_NAME
                    right_run.font.size = Pt(FONT_SIZE_SECTION_CONTENT)
                for bullet in exp.bullets:
                    p = doc.add_paragraph()
                    p.style = "ResumeBullet"
                    p.paragraph_format.left_indent = Inches(BULLET_INDENT)
                    p.paragraph_format.first_line_indent = Inches(-BULLET_INDENT)
                    run = p.add_run("• " + bullet)
                    run.font.name = FONT_NAME
                    run.font.size = Pt(FONT_SIZE_BULLET)
                doc.add_paragraph()
            last_para = doc.paragraphs[-1]
            last_para.paragraph_format.space_after = Pt(SECTION_SPACING)
        if resume.projects:
            DocxRenderer._render_section_title(doc, "PROJECTS")
            for proj in resume.projects:
                title_line = proj.title
                if proj.tech:
                    title_line += f" | {proj.tech}"
                p = doc.add_paragraph(title_line)
                p.style = "ResumeHeading"
                p.runs[0].bold = True
                p.runs[0].font.name = FONT_NAME
                p.runs[0].font.size = Pt(FONT_SIZE_SECTION_CONTENT)
                for bullet in proj.bullets:
                    p = doc.add_paragraph()
                    p.style = "ResumeBullet"
                    p.paragraph_format.left_indent = Inches(BULLET_INDENT)
                    p.paragraph_format.first_line_indent = Inches(-BULLET_INDENT)
                    run = p.add_run("• " + bullet)
                    run.font.name = FONT_NAME
                    run.font.size = Pt(FONT_SIZE_BULLET)
                doc.add_paragraph()
            last_para = doc.paragraphs[-1]
            last_para.paragraph_format.space_after = Pt(SECTION_SPACING)
        if resume.education:
            DocxRenderer._render_section_title(doc, "EDUCATION")
            for edu in resume.education:
                edu_line = f"{edu.degree} – {edu.institution}"
                p = doc.add_paragraph()
                p.style = "ResumeHeading"
                left_run = p.add_run(edu_line)
                left_run.bold = True
                left_run.font.name = FONT_NAME
                left_run.font.size = Pt(FONT_SIZE_SECTION_CONTENT)
                if edu.year:
                    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.0))
                    p.add_run("\t").bold = False
                    right_run = p.add_run(edu.year)
                    right_run.font.name = FONT_NAME
                    right_run.font.size = Pt(FONT_SIZE_SECTION_CONTENT)
                for detail in edu.details:
                    p = doc.add_paragraph()
                    p.style = "ResumeBullet"
                    p.paragraph_format.left_indent = Inches(BULLET_INDENT)
                    p.paragraph_format.first_line_indent = Inches(-BULLET_INDENT)
                    if detail.startswith("GPA:") or detail.startswith("Relevant Coursework:") or detail.startswith("Awards:"):
                        run = p.add_run(detail)
                    else:
                        run = p.add_run("• " + detail)
                    run.font.name = FONT_NAME
                    run.font.size = Pt(FONT_SIZE_BULLET)
                doc.add_paragraph()
            last_para = doc.paragraphs[-1]
            last_para.paragraph_format.space_after = Pt(SECTION_SPACING)
        if (resume.skills.programming_languages or resume.skills.frameworks_libraries or
            resume.skills.tools_technologies or resume.skills.methodologies):
            DocxRenderer._render_section_title(doc, "TECHNICAL SKILLS")
            if resume.skills.programming_languages:
                p = doc.add_paragraph()
                p.style = "ResumeNormal"
                run = p.add_run(f"Programming Languages: {', '.join(resume.skills.programming_languages)}")
                run.font.name = FONT_NAME
                run.font.size = Pt(FONT_SIZE_SECTION_CONTENT)
            if resume.skills.frameworks_libraries:
                p = doc.add_paragraph()
                p.style = "ResumeNormal"
                run = p.add_run(f"Frameworks & Libraries: {', '.join(resume.skills.frameworks_libraries)}")
                run.font.name = FONT_NAME
                run.font.size = Pt(FONT_SIZE_SECTION_CONTENT)
            if resume.skills.tools_technologies:
                p = doc.add_paragraph()
                p.style = "ResumeNormal"
                run = p.add_run(f"Tools & Technologies: {', '.join(resume.skills.tools_technologies)}")
                run.font.name = FONT_NAME
                run.font.size = Pt(FONT_SIZE_SECTION_CONTENT)
            if resume.skills.methodologies:
                p = doc.add_paragraph()
                p.style = "ResumeNormal"
                run = p.add_run(f"Methodologies: {', '.join(resume.skills.methodologies)}")
                run.font.name = FONT_NAME
                run.font.size = Pt(FONT_SIZE_SECTION_CONTENT)
            last_para = doc.paragraphs[-1]
            last_para.paragraph_format.space_after = Pt(SECTION_SPACING)
        if resume.certifications:
            DocxRenderer._render_section_title(doc, "CERTIFICATIONS")
            for cert in resume.certifications:
                p = doc.add_paragraph()
                p.style = "ResumeBullet"
                p.paragraph_format.left_indent = Inches(BULLET_INDENT)
                p.paragraph_format.first_line_indent = Inches(-BULLET_INDENT)
                run = p.add_run("• " + cert)
                run.font.name = FONT_NAME
                run.font.size = Pt(FONT_SIZE_BULLET)
            last_para = doc.paragraphs[-1]
            last_para.paragraph_format.space_after = Pt(SECTION_SPACING)
        if resume.hobbies:
            DocxRenderer._render_section_title(doc, "HOBBIES")
            p = doc.add_paragraph(", ".join(resume.hobbies))
            p.style = "ResumeNormal"
            p.paragraph_format.space_after = Pt(SECTION_SPACING)
        if resume.career_goal:
            DocxRenderer._render_section_title(doc, "CAREER TARGET")
            p = doc.add_paragraph(resume.career_goal)
            p.style = "ResumeNormal"
        doc.save(output_path)
    @staticmethod
    def _setup_styles(doc):
        style = doc.styles.add_style('ResumeNormal', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = FONT_NAME
        style.font.size = Pt(FONT_SIZE_SECTION_CONTENT)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(4)
        style = doc.styles.add_style('ResumeHeading', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = FONT_NAME
        style.font.size = Pt(FONT_SIZE_SECTION_CONTENT)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        style.paragraph_format.space_before = Pt(2)
        style.paragraph_format.space_after = Pt(2)
        style = doc.styles.add_style('ResumeBullet', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = FONT_NAME
        style.font.size = Pt(FONT_SIZE_BULLET)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.left_indent = Inches(BULLET_INDENT)
        style.paragraph_format.first_line_indent = Inches(-BULLET_INDENT)
    @staticmethod
    def _render_header(doc, header):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header.name.upper())
        run.bold = True
        run.font.name = FONT_NAME
        run.font.size = Pt(FONT_SIZE_HEADER_NAME)
        p.paragraph_format.space_after = Pt(6)
        contact_parts = []
        if header.address:
            contact_parts.append(header.address)
        if header.phone:
            contact_parts.append(header.phone)
        if header.email:
            contact_parts.append(header.email)
        if header.linkedin:
            contact_parts.append(header.linkedin)
        if header.github:
            contact_parts.append(header.github)
        if contact_parts:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(" | ".join(contact_parts))
            run.font.name = FONT_NAME
            run.font.size = Pt(FONT_SIZE_HEADER_CONTACT)
            p.paragraph_format.space_after = Pt(12)
    @staticmethod
    def _render_section_title(doc, title_text):
        p = doc.add_paragraph()
        run = p.add_run(title_text)
        run.bold = True
        run.font.name = FONT_NAME
        run.font.size = Pt(FONT_SIZE_SECTION_TITLE)
        run.underline = True
        p.paragraph_format.space_before = Pt(SECTION_SPACING)
        p.paragraph_format.space_after = Pt(6)
    @staticmethod
    def _add_bullet_paragraph(doc, text, indent_level=0):
        p = doc.add_paragraph()
        indent = Inches(BULLET_INDENT * (indent_level + 1))
        p.paragraph_format.left_indent = indent
        p.paragraph_format.first_line_indent = Inches(-BULLET_INDENT)
        run = p.add_run("• " + text)
        run.font.name = FONT_NAME
        run.font.size = Pt(FONT_SIZE_BULLET)
        return p